import nltk
from nltk.corpus import wordnet
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
import spacy
from docx import Document
import pickle
import os

from . import keyword_extraction
from . import pipeline 

class Doc:
    def __init__(self, text, keywords):
        self.text = text
        self.keywords = keywords

# getter and setter for inputs
class inputs:
    def __init__(self, file_path, query: str):
        self.file_path = file_path
        self.query = query

    def __init__(self):
        return

    def set_file_path(self, file_path):
        self.file_path = file_path

    def set_query(self, query):
        self.query = query

stopwords = set(stopwords.words('english'))
nlp = spacy.load('en_core_web_lg') # py -m spacy download en_core_web_lg
lemmatizer = WordNetLemmatizer()

def preprocess_text(text: str) -> str:
    ''' 
    remove extra spaces, newlines, tabs, stopwords, numbers
    we will end up with a list of unique and significant words
    '''
    text = nltk.re.sub(r'\s+', ' ', text).strip()
    text = text.replace('\n', '').replace('\t', '')
    text = word_tokenize(text)
    text = [lemmatizer.lemmatize(word.lower()) for word in text if word not in stopwords and word.isalpha()]
    text = set(text)

    #print(text)
    #return a single text string
    return ' '.join(list(text))

# use spacy to compute similarity between query and each document
def get_matching_documents(query: str, documents: list) -> list:
    #print('Calculating similarities...')
    
    error_threshold = 0.5 # similarity score threshold; subject to change
    
    #print(f'matching func|| query: {query} , documents: {documents}')
    
    query = nlp(preprocess_text(query))
    nlp_docs = []
    for document in documents:
        nlp_docs.append(nlp(preprocess_text(document)))
    
    similarity_scores = [query_similarity := query.similarity(doc) for doc in nlp_docs]
    # print(similarity_scores)
    
    # map documents to their similarities
    doc_similarity_tuples = []
    for i, doc in enumerate(documents):
        doc_similarity_tuples.append((doc, similarity_scores[i]))
    
    # descending sort
    sorted_docs_by_score = sorted(doc_similarity_tuples, key=lambda x:x[1], reverse=True)
    # print(sorted_docs_by_score)
    
    document_matches = []
    for doc, score in sorted_docs_by_score:
        if score > error_threshold:
            if doc not in document_matches:
                # match found
                print(f'Score: {score}')
                document_matches.append(doc)
    #print('Done calculating similarities.')
    
    return document_matches
    
def extract_text_from_tables(docx_path: str) -> list:
    '''
    extract text from tables in a docx by row
    everything within the same row is put into the same document
    '''
    #print('Extracting text from tables...')
    
    docx = Document(docx_path)
    table_documents = []
    for table in docx.tables:
        for row in table.rows:
            row_content = ''
            for cell in row.cells:
                row_content += cell.text + ' '
            table_documents.append(row_content)
    
    #print('Done extracting text from tables.')
    
    return table_documents
    
def extract_text_from_bold(docx_path: str) -> list:
    '''
    extract text using the position of bolded text
    everything after the bolded text is added to the same document
    until the next bolded text
    '''
    docx = Document(docx_path)
    bolded_documents = []
    
    counter = 0
    for paragraph in docx.paragraphs:
        for run in paragraph.runs:
            if run.bold:
                bolded_documents.append(run.text)
                counter += 1
            elif counter == 0: # whenever the docx starts with a non-bold run
                bolded_documents.append(run.text)
            else:
                bolded_documents[counter] += run.text
    
    return bolded_documents  

def process_course_material(file_path, kw_model):
    print('Processing course material...')
    with open(file_path, 'r') as f:
        text = f.read()
        keywords = keyword_extraction.extract_keywords(kw_model, text, ngram_range=3, top_n=10)
        #print(f'Course material keywords: {keywords}')
        doc = [Doc(text, keywords)]
        #print(doc)
        with open(pickled_name(file_path), 'wb') as p:
            pickle.dump(doc, p)

# TENTATIVE
# How and where we save the keywords of a document may change (i.e. in database instead)
def process_course_syllabus(file_path, kw_model):
    '''
    Pull text in tables and organize them into a list of 'documents'
    for each document, determine keywords for each document and save them
    to a class containing both the text and keywords
    finally, pickle a list of these course_material document classes for later use
    '''
    print('Processing syllabus...')
    
    table_documents = extract_text_from_tables(file_path)
    
    docs = []
    for document in table_documents:
        keywords = keyword_extraction.extract_keywords(kw_model, document, ngram_range=3, top_n=10)
        #print(keywords)
        docs.append(Doc(document, keywords))
    
    with open(pickled_name(file_path), 'wb') as f:
        pickle.dump(docs, f)
    
    #print('Done processing course_material.')

# Checks if file has already been processed
def process_file(file_path, kw_model):
    if not os.path.exists(pickled_name(file_path)):
        if 'docx' in file_path:              
            process_course_syllabus(file_path, kw_model)
            return
        process_course_material(file_path, kw_model)
        
# Changes the file name to .pickle
def pickled_name(file_path) -> str:
    dot_index = file_path.rfind(".")
    pickled_path = file_path[:dot_index] + ".pickle"
    return pickled_path

# 
def ask_bert(file_path: str, query: str):
    new_query = inputs()
    new_query.set_file_path(file_path)
    new_query.set_query(query)

    ask_bert_detailed(new_query)

# Takes in question and file and outputs answer with confidence score
def ask_bert_detailed(new_query: inputs):

    kw_model = keyword_extraction.load()                    

    process_file(new_query.file_path, kw_model)
    only_key_docs = []

    with open(pickled_name(new_query.file_path), 'rb') as f:
        
        all_docs = pickle.load(f)
        #print(all_docs)
        # only retrieve the documents that contain the keywords
        query_keywords = keyword_extraction.extract_keywords(kw_model, new_query.query, ngram_range=2, top_n=3)
        for doc in all_docs:
            doc_keywords = ' '.join(doc.keywords)
            for keyword in query_keywords:
                if keyword in doc_keywords and doc_keywords not in only_key_docs:
                    only_key_docs.append(doc_keywords)
                    print(f'Added doc keywords: {doc_keywords}')


    # map keywords to documents
    doc_dict = {}
    for doc in all_docs:
         doc_dict[' '.join(doc.keywords)] = doc.text

    print(f'Transformed query: {query_keywords}') 
    matches = get_matching_documents(' '.join(query_keywords), only_key_docs)
    # retrieves top 5 matches
    context = ""
    for i, match in enumerate(matches[:5]):
        context += (doc_dict[match] + ' ')
    
    print("Context:\n" + context + '\n\n')
    pipeline = Pipeline()
    try:
        pipeline.get_answer(new_query.query, context)
    except ValueError:
        print('Error: No context or query given')

if __name__ == '__main__':
    ask_bert("a2-3377.txt", "10 frequently used")
    ask_bert("Syllabus-3377-converted.docx", "What is the grading policy?")