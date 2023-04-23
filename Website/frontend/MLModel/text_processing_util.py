from docx import Document
import pickle
import os
import mysql.connector
from decouple import config

import keyword_extraction
from pipeline import *

class Doc:
    def __init__(self, text: str, keywords: list, origin: str | int):
        self.text = text # full text of one theme
        self.keywords = keywords # words that highlight the document's theme
        self.origin = origin # any course material: str | thread ID: int
        
    def __init__(self):
        return
        
    def set_text(self, text: str):
        self.text = text
    
    def set_keywords(self, keywords: list):
        self.keywords = keywords
        
    def set_origin(self, origin: str | int):
        self.origin = origin

        print(f'New Doc obj origin type: {type(origin)}')
        if type(origin) is str:
            self.type = 'file'
        else:
            self.type =  'thread'
    
    


def extract_text_from_tables(docx_path: str) -> list:
    '''
    extract text from tables in a docx by row
    everything within the same row is put into the same document
    '''
    #print('Extracting text from tables...')
    
    docx = Document(docx_path)
    table_documents = []
    for table in docx.tables:
        for row in table.rows:
            row_content = ''
            for cell in row.cells:
                row_content += cell.text + ' '
            table_documents.append(row_content)
    
    #print('Done extracting text from tables.')
    
    return table_documents
    
def extract_text_from_bold(docx_path: str) -> list:
    '''
    extract text using the position of bolded text
    everything after the bolded text is added to the same document
    until the next bolded text
    '''
    docx = Document(docx_path)
    bolded_documents = []
    
    counter = 0
    for paragraph in docx.paragraphs:
        for run in paragraph.runs:
            if run.bold:
                bolded_documents.append(run.text)
                counter += 1
            elif counter == 0: # whenever the docx starts with a non-bold run
                bolded_documents.append(run.text)
            else:
                bolded_documents[counter] += run.text
    
    return bolded_documents  

def existing_doc(class_data: list, current_file_origin: int | str) -> int:
    
    for doc in class_data:
        if doc.origin == current_file_origin:
            return class_data.index(doc)
    return -1

def process_course_material(class_data_path, file_path, kw_model):

    try:
        with open(class_data_path, 'rb') as p:
            class_data = pickle.load(p) # list of processed Doc objs
    except FileNotFoundError:
        class_data = []
        
    with open(file_path, 'r') as f:
        
        text = f.read()
        keywords = keyword_extraction.extract_keywords(kw_model, text, ngram_range=3, top_n=10)
        
        existing_doc_index = existing_doc(class_data, file_path)
        
        if existing_doc_index != -1: # just update the keywords
            class_data[existing_doc_index].keywords = keywords

        else:
            doc = Doc()
            doc.set_text(text)
            doc.set_keywords(keywords)
            doc.set_origin(str(file_path))
            
            class_data.append(doc)
            
        with open(class_data_path, 'wb') as new_p:
            pickle.dump(class_data, new_p) # overwrite

# TENTATIVE
# How and where we save the keywords of a document may change (i.e. in database instead)
def process_course_syllabus(class_data_path, file_path, kw_model):
    '''
    Pull text in tables and organize them into a list of 'documents'
    for each document, determine keywords for each document and save them
    to a class containing both the text and keywords
    finally, pickle a list of these course_material document classes for later use
    '''
    print('Processing syllabus...')
    
    table_documents = extract_text_from_tables(file_path)
    try:
        with open(class_data_path, 'rb') as p:
            class_data = pickle.load(p) # list of processed Doc objs
    except FileNotFoundError:
        class_data = []

    existing_doc_start_index = existing_doc(class_data, file_path)
    print(f'Existing index start: {existing_doc_start_index}')
    
    for i, document in enumerate(table_documents):
        keywords = keyword_extraction.extract_keywords(kw_model, document, ngram_range=3, top_n=10)


        if existing_doc_start_index != -1: # just update the keywords
            class_data[existing_doc_start_index + i].keywords = keywords
            
        else:
            doc = Doc()
            doc.set_text(document)
            doc.set_keywords(keywords)
            doc.set_origin(str(file_path))
            
            class_data.append(doc)
    

    print(f'Updated class data: {class_data}')
    print('Length: ', len(class_data))
    with open(class_data_path, 'wb') as new_p:
        pickle.dump(class_data, new_p) # overwrite
    
    #print('Done processing course_material.')
    
'''
# Changes the file name to .pickle
def pickled_name(file_path) -> str:
    dot_index = file_path.rfind(".")
    pickled_path = file_path[:dot_index] + ".pickle"
    return pickled_path
'''

# Checks if file has already been processed
def process_file(class_data_path, file_path, kw_model):
    if '.docx' in file_path:  
        print('Recognized as .docx')            
        process_course_syllabus(class_data_path, file_path, kw_model)
        return
    # ends with .txt
    print('Recognized as .txt')
    process_course_material(class_data_path, file_path, kw_model)
        

def process_thread(class_data_path, thread_ID, kw_model):
    threadtext = get_thread_and_reply_text(thread_ID)
    
    try:
        with open(class_data_path, 'rb') as p:
            class_data = pickle.load(p) # list of processed Doc objs
    except FileNotFoundError:
        class_data = []

    keywords = keyword_extraction.extract_keywords(kw_model, threadtext, ngram_range=3, top_n=10)
    
    existing_doc_index = existing_doc(class_data, threadtext)
    
    if existing_doc_index != -1: # just update the keywords
        class_data[existing_doc_index].keywords = keywords
    else:
        doc = Doc()
        doc.set_text(threadtext)
        doc.set_keywords(keywords)
        doc.set_origin(thread_ID)
        
        class_data.append(doc)
        
    with open(class_data_path, 'wb') as new_p:
        pickle.dump(class_data, new_p) # overwrite



def get_thread_and_reply_text(thread_ID):
    
    db = mysql.connector.connect(
        host = '10.176.67.73',
        port = '3306',
        user = 'db_user',
        password = config("DB_CREDS")
    )
    
    dbcursor = db.cursor()

    dbcursor.execute("SELECT DISTINCT t.ThreadTitle, t.ThreadContent, r.Content FROM Thread as t, Reply as r WHERE t.ThreadID = (%s) AND t.ThreadID = r.Thread_ThreadID", (thread_ID))

    result = dbcursor.fetchall() # Grab query result from cursor
    
    dbcursor.close()
    
    text = ''
    for x in result:
        text += x
    return text

    
    

def new_process(class_ID: int, thread_ID, file_path = ''):
    # Process a thread or file
    class_data_path = 'class_data/' + str(class_ID) + '.pickle' #TENTATIVE (e.g. 12345.pickle)
    print(f'Class data path: {class_data_path}')
    #print('Processing...')
    kw_model = keyword_extraction.load()
    
    if thread_ID != -1:
        print('Processing a thread')
        process_thread(class_data_path, thread_ID, kw_model)
    if file_path != '':
        print('Processing a file')
        process_file(class_data_path, file_path, kw_model)
    #print('Done processing.')
    

def del_document(class_ID: int, origin: int | str):
    
    class_data_path = 'class_data/' + str(class_ID) + '.pickle'

    with open(class_data_path, 'rb') as o:
        class_data = pickle.load(o)
        
    class_data = [doc for doc in class_data if doc.origin != origin]
    
    print(f'Updated class data: {class_data}')
    print('Length: ', len(class_data))
    with open(class_data_path, 'wb') as c:
        pickle.dump(class_data, c) # overwrite

if __name__ == '__main__':

    #new_process(12345, file_path = 'Syllabus-3377-converted.docx')
    new_process(class_id = 12345, thread_ID=6)

    #del_document(12345, 'Syllabus-3377-converted.docx')