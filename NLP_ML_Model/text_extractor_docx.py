import nltk
from nltk.corpus import wordnet
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
import spacy
from docx import Document
import pickle
import os

import keyword_extraction
from pipeline import *

class Doc:
    def __init__(self, text, keywords):
        self.text = text
        self.keywords = keywords

# getter and setter for inputs
class inputs:
    def __init__(self, docx_path, query):
        self.docx_path = docx_path
        self.query = query

    def __init__(self):
        return

    def set_docx_path(self, docx_path):
        self.docx_path = docx_path

    def set_query(self, query):
        self.query = query

stopwords = set(stopwords.words('english'))
nlp = spacy.load('en_core_web_lg') # py -m spacy download en_core_web_lg
lemmatizer = WordNetLemmatizer()

def preprocess_text(text: str) -> str:
    ''' 
    remove extra spaces, newlines, tabs, stopwords, numbers
    we will end up with a list of unique and significant words
    '''
    text = nltk.re.sub(r'\s+', ' ', text).strip()
    text = text.replace('\n', '').replace('\t', '')
    text = word_tokenize(text)
    text = [lemmatizer.lemmatize(word.lower()) for word in text if word not in stopwords and word.isalpha()]
    text = set(text)

    #print(text)
    #return a single text string
    return ' '.join(list(text))

# use spacy to compute similarity between query and each document
def get_matching_documents(query: str, documents: list) -> list:
    #print('Calculating similarities...')
    
    error_threshold = 0.5 # similarity score threshold; subject to change
    
    #print(f'matching func|| query: {query} , documents: {documents}')
    
    query = nlp(preprocess_text(query))
    nlp_docs = []
    for document in documents:
        nlp_docs.append(nlp(preprocess_text(document)))
    
    similarity_scores = [query_similarity := query.similarity(doc) for doc in nlp_docs]
    # print(similarity_scores)
    
    # map documents to their similarities
    doc_similarity_tuples = []
    for i, doc in enumerate(documents):
        doc_similarity_tuples.append((doc, similarity_scores[i]))
    
    # descending sort
    sorted_docs_by_score = sorted(doc_similarity_tuples, key=lambda x:x[1], reverse=True)
    # print(sorted_docs_by_score)
    
    document_matches = []
    for doc, score in sorted_docs_by_score:
        if score > error_threshold:
            if doc not in document_matches:
                # match found
                print(f'Score: {score}')
                document_matches.append(doc)
    #print('Done calculating similarities.')
    
    return document_matches
    
def extract_text_from_tables(docx_path: str) -> list:
    '''
    extract text from tables in a docx by row
    everything within the same row is put into the same document
    '''
    #print('Extracting text from tables...')
    
    docx = Document(docx_path)
    table_documents = []
    for table in docx.tables:
        for row in table.rows:
            row_content = ''
            for cell in row.cells:
                row_content += cell.text + ' '
            table_documents.append(row_content)
    
    #print('Done extracting text from tables.')
    
    return table_documents
    
def extract_text_from_bold(docx_path: str) -> list:
    '''
    extract text using the position of bolded text
    everything after the bolded text is added to the same document
    until the next bolded text
    '''
    docx = Document(docx_path)
    bolded_documents = []
    
    counter = 0
    for paragraph in docx.paragraphs:
        for run in paragraph.runs:
            if run.bold:
                bolded_documents.append(run.text)
                counter += 1
            elif counter == 0: # whenever the docx starts with a non-bold run
                bolded_documents.append(run.text)
            else:
                bolded_documents[counter] += run.text
    
    return bolded_documents  

# TENTATIVE
# How and where we save the keywords of a document may change (i.e. in database instead)
def process_course_material(docx_path, kw_model):
    '''
    Pull text in tables and organize them into a list of 'documents'
    for each document, determine keywords for each document and save them
    to a class containing both the text and keywords
    finally, pickle a list of these course_material document classes for later use
    '''
    print('Processing course_material...')
    
    table_documents = extract_text_from_tables(docx_path)
    
    docs = []
    for document in table_documents:
        keywords = keyword_extraction.extract_keywords(kw_model, document, ngram_range=3, top_n=10)
        #print(keywords)
        docs.append(Doc(document, keywords))
    
    with open('course_material_docs.pickle', 'wb') as f:
        pickle.dump(docs, f)
    
    #print('Done processing course_material.')

# not really sure if we need this tbh
def get_similar_words(keyword):
    similar_words = []
    
    for syn in wordnet.synsets(keyword):
        for synonym in syn.lemmas():
            similar_words.append(synonym.name())
    print(f'Similar words: {set(similar_words)}')

# Checks if docx has already been processed
def process_docx(docx_path, kw_model):
    if not os.path.exists('course_material_docs.pickle'):               #TODO: Make pickle file names dynamic to given docx name
        process_course_material(docx_path, kw_model)

def ask_bert(docx_path, query):
    new_query = inputs()
    new_query.set_docx_path(docx_path)
    new_query.set_query(query)

    ask_bert_detailed(new_query)

# Takes in question and docx and outputs answer with confidence score
def ask_bert_detailed(new_query):

    kw_model = keyword_extraction.load()                                #TODO: fix kw_model needs

    process_docx(new_query.docx_path, kw_model)

    with open('course_material_docs.pickle', 'rb') as f:
        all_docs = pickle.load(f)
        
        # only retrieve the documents that contain the keywords
        only_key_docs = []
        query_keywords = keyword_extraction.extract_keywords(kw_model, new_query.query, ngram_range=2, top_n=3)
        for doc in all_docs:
            for keyword in query_keywords:
                doc_keywords = ' '.join(doc.keywords)
                if keyword in doc_keywords:
                    only_key_docs.append(doc_keywords)
                    #print(f'Added doc keywords: {doc_keywords}')

    # map keywords to documents
    doc_dict = {}
    for doc in all_docs:
         doc_dict[' '.join(doc.keywords)] = doc.text

    print(f'Transformed query: {query_keywords}') 
    matches = get_matching_documents(' '.join(query_keywords), only_key_docs)
    # retrieves top 5 matches
    context = ""
    for i, match in enumerate(matches[:5]):
        context += (doc_dict[match] + ' ')
    
    # check if there are no matches
    if context.isspace():
        raise Exception("There is no context available.")
    
    print("Context:\n" + context + '\n\n')
    pipeline = Pipeline()
    pipeline.get_answer(new_query.query, context)

if __name__ == '__main__':
    ask_bert("Syllabus-3377-converted.docx", "Who is the professor?")